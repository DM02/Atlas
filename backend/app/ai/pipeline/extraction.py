import io
from dataclasses import dataclass

import docx
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
from pypdf import PdfReader

MIME_PDF = "application/pdf"
MIME_TXT = "text/plain"
MIME_DOCX = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

EXTENSION_TO_MIME = {
    ".pdf": MIME_PDF,
    ".txt": MIME_TXT,
    ".docx": MIME_DOCX,
}


@dataclass
class ExtractedPage:
    page_number: int | None
    text: str


def extract_text(content: bytes, mime_type: str) -> list[ExtractedPage]:
    if mime_type == MIME_PDF:
        return _extract_pdf(content)
    if mime_type == MIME_TXT:
        return [ExtractedPage(page_number=None, text=content.decode("utf-8", errors="replace"))]
    if mime_type == MIME_DOCX:
        return _extract_docx(content)
    raise ValueError(f"Unsupported mime type: {mime_type}")


def _extract_pdf(content: bytes) -> list[ExtractedPage]:
    reader = PdfReader(io.BytesIO(content))
    return [
        ExtractedPage(page_number=i + 1, text=page.extract_text() or "")
        for i, page in enumerate(reader.pages)
    ]


def _extract_docx(content: bytes) -> list[ExtractedPage]:
    document = docx.Document(io.BytesIO(content))
    parts = [_block_text(block) for block in _iter_block_items(document)]
    text = "\n".join(part for part in parts if part)
    return [ExtractedPage(page_number=None, text=text)]


def _iter_block_items(document: docx.Document):
    """Yields paragraphs and tables in the order they appear in the document
    body. python-docx's own `.paragraphs`/`.tables` each return a flat list of
    just that one type, losing how they're interleaved in the actual document
    — this walks the underlying XML body directly instead (the standard
    recipe for reading a docx in original document order), so a table stays
    next to the paragraphs around it rather than being silently dropped
    entirely. Found via a real .docx with 7 tables (including one holding the
    OSI layer names/PDU/functions) that a plain `.paragraphs` join skipped
    completely — every table in the document was invisible to retrieval.
    """
    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, document)


def _block_text(block: Paragraph | Table) -> str:
    if isinstance(block, Table):
        return "\n".join(" | ".join(cell.text for cell in row.cells) for row in block.rows)
    return block.text
