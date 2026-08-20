import io

import docx
import pytest

from app.ai.pipeline.extraction import MIME_DOCX, MIME_TXT, extract_text


def test_extract_text_txt_returns_single_page() -> None:
    pages = extract_text(b"hello atlas", MIME_TXT)

    assert len(pages) == 1
    assert pages[0].page_number is None
    assert pages[0].text == "hello atlas"


def test_extract_text_docx_joins_paragraphs() -> None:
    document = docx.Document()
    document.add_paragraph("First paragraph.")
    document.add_paragraph("Second paragraph.")
    buffer = io.BytesIO()
    document.save(buffer)

    pages = extract_text(buffer.getvalue(), MIME_DOCX)

    assert len(pages) == 1
    assert "First paragraph." in pages[0].text
    assert "Second paragraph." in pages[0].text


def test_extract_text_rejects_unsupported_mime_type() -> None:
    with pytest.raises(ValueError, match="Unsupported mime type"):
        extract_text(b"data", "application/zip")


def test_extract_text_docx_includes_table_content() -> None:
    # Real bug found via live testing: a plain `.paragraphs` join silently
    # drops every table in the document — this reproduces that with a
    # minimal table so it can't regress unnoticed again.
    document = docx.Document()
    document.add_paragraph("Intro paragraph.")
    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Nr"
    table.rows[0].cells[1].text = "Warstwa"
    table.rows[1].cells[0].text = "7"
    table.rows[1].cells[1].text = "Aplikacji"
    buffer = io.BytesIO()
    document.save(buffer)

    pages = extract_text(buffer.getvalue(), MIME_DOCX)

    assert "Intro paragraph." in pages[0].text
    assert "Warstwa" in pages[0].text
    assert "Aplikacji" in pages[0].text


def test_extract_text_docx_preserves_paragraph_and_table_order() -> None:
    document = docx.Document()
    document.add_paragraph("Before table.")
    table = document.add_table(rows=1, cols=1)
    table.rows[0].cells[0].text = "Inside table."
    document.add_paragraph("After table.")
    buffer = io.BytesIO()
    document.save(buffer)

    pages = extract_text(buffer.getvalue(), MIME_DOCX)

    text = pages[0].text
    assert text.index("Before table.") < text.index("Inside table.") < text.index("After table.")
